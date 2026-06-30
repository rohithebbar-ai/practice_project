#!/usr/bin/env python3
"""
generate_doc.py
Generates BOTH a .docx and a .pdf of the Indent Validation Service
architecture document.

Requirements:
    pip install python-docx reportlab

Run:
    python generate_doc.py

Output (written next to this script):
    Indent_Validation_Service_Architecture.docx
    Indent_Validation_Service_Architecture.pdf
"""

import os

# ─────────────────────────────────────────────────────────────────────────────
# CONTENT  (edit text here if needed)
# ─────────────────────────────────────────────────────────────────────────────

TITLE    = "Indent Validation Service"
SUBTITLE = "Production Architecture, Integration & Data Flow"
VERSION  = "Version 1.1 (Draft)"

# Diagrams are kept as monospace ASCII blocks (rendered in a fixed-width font
# in both outputs so the boxes line up).

ARCH_DIAGRAM = r"""
                          +-------------------------+
                          |      IPMS  PORTAL       |
                          | 'Check Adequacy' click  |
                          |    displays results     |
                          +------------+------------+
                                       |  indent_id
                                       v
  +-----------------+       +-------------------------+       +------------------+
  | IPMS FIELD API  |<------|                         |------>|  DOCUMENT STORE  |
  | indent fields   |       |   INDENT VALIDATION     |       | attachments      |
  +-----------------+       |        SERVICE          |       | (OAuth)          |
                            |   (stateless compute)   |       +------------------+
  +-----------------+       |  1 fetch fields + docs  |
  |   GENAI  API    |<------|  2 classify by content  |
  | authenticated   |       |  3 infer + analyse      |
  | extraction      |       |  4 score + return JSON  |
  +-----------------+       +------------+------------+
                                       |  scored analysis
                                       v
                          +-------------------------+
                          |     IPMS  DATABASE      |
                          | score + full analysis   |
                          | history + feedback      |
                          +------------+------------+
                                       ^
                                       |  publishes new standard version
                          +------------+------------+
                          | OFFLINE STANDARD        |
                          | PIPELINE (scheduled)    |
                          | rebuilds the standard   |
                          +-------------------------+
"""

FLOW_STEPS = [
    ("1.  User clicks \u201CCheck Adequacy\u201D in IPMS",
     ["Advisory action \u2014 never blocks submission"]),
    ("2.  IPMS calls the Validation API (async)",
     ["POST /check-adequacy { indent_id, domain }",
      "Returns immediately: { analysis_id, status }"]),
    ("3.  Fetch fields and attachments (in parallel)",
     ["3a. Indent fields from IPMS Field API: [{question, answer, required}, ...]",
      "3b. All attachments from the document store (any filenames)"]),
    ("4.  Classify every document by content",
     ["Filenames and upload slots are not trusted",
      "Each file is parsed and identified by what it contains",
      "e.g. \u201CTSK_name.pdf\u201D is detected as a BOQ; a drawing inside a spec is found",
      "Builds a true inventory of the documents that exist"]),
    ("5.  Infer requirements",
     ["From the field choices + domain:",
      "Type = Service        -> certain documents expected",
      "Single party = Yes    -> party justification expected",
      "High risk             -> HSE plan expected",
      "Produces: this indent should contain X, Y, Z"]),
    ("6.  Analyse via authenticated GENAI API",
     ["Send: field data + classified document content + inferred requirements + standard",
      "Get: extraction, interrelationships, gaps"]),
    ("7.  Compare and score (advisory)",
     ["What is present vs what was inferred as needed",
      "Interrelationship checks (fields vs BOQ, scope vs items, risk vs HSE)",
      "Produces: score + grade + findings + missing + recommendations"]),
    ("8.  Return full analysis to IPMS",
     ["{ indent_id, analysis_timestamp, score, grade, findings,",
      "  recommendations, missing_documents, interrelationship_issues }",
      "The service does not write to the database"]),
    ("9.  IPMS persists the analysis",
     ["Stores indent_id, timestamp, score AND full detailed analysis",
      "Versioning + history maintained by IPMS",
      "Always shows the latest only \u2014 to user and manager"]),
    ("10.  Iterate",
     ["User sees full analysis -> improves -> re-checks",
      "Submits when satisfied (at any score)",
      "Manager sees the same full analysis (latest only)",
      "Manager may send back -> user improves -> resubmit"]),
]

# Tables: (title, headers, rows, col_widths_for_pdf)
APIS_REQUIRED = (
    ["API", "Direction", "Purpose"],
    [
        ["Check-Adequacy API", "IPMS to Service (exposed by us)", "Receives the request, returns the scored analysis"],
        ["Indent Field API",   "Service to IPMS (consumed)",      "Provides the indent's filled field data"],
        ["Document Store API", "Service to Store (consumed)",     "Provides the indent's attachments via OAuth"],
        ["GENAI API",          "Service to GENAI (consumed)",     "Performs extraction and analysis (authenticated)"],
    ],
)

REQ_REQUEST = (
    ["Key", "Type", "Description"],
    [
        ["indent_id",    "string", "Identifier of the indent to validate"],
        ["domain",       "string", "Indent domain (selects the standard)"],
        ["callback_url", "string", "Where the result is delivered when ready (async)"],
    ],
)

REQ_ACK = (
    ["Key", "Type", "Description"],
    [
        ["analysis_id", "string", "Identifier for this analysis run"],
        ["status",      "string", "Processing state (e.g. processing)"],
    ],
)

RESULT_TBL = (
    ["Key", "Type", "Description"],
    [
        ["analysis_id",              "string",          "Identifier for this analysis run"],
        ["indent_id",                "string",          "The indent that was analysed"],
        ["domain",                   "string",          "Indent domain"],
        ["analysis_timestamp",       "string",          "When the analysis ran (ISO timestamp)"],
        ["standard_version",         "string",          "Which standard version was used"],
        ["score",                    "number",          "Overall adequacy score, 0-100"],
        ["grade",                    "string",          "Strong / Adequate / Needs Improvement / Weak"],
        ["score_breakdown",          "object",          "Points per category"],
        ["findings",                 "list of objects", "Each: category, status, title, detail"],
        ["recommendations",          "list of strings", "Actions to improve the indent"],
        ["missing_documents",        "list of strings", "Required document types not found"],
        ["interrelationship_issues", "list of strings", "Cross-document inconsistencies"],
        ["document_inventory",       "list of objects", "Each: file, classified_as, genuine (boolean)"],
    ],
)

FIELD_API_TBL = (
    ["Key (per field object)", "Type", "Description"],
    [
        ["question", "string",  "The field label shown to the user"],
        ["answer",   "string",  "The value selected or entered"],
        ["required", "boolean", "Whether the field was mandatory (the * fields)"],
    ],
)

TABLES_REQUIRED = (
    ["Table", "Purpose"],
    [
        ["indent_analysis",     "Stores every analysis run - score plus full detailed analysis; history with a latest flag"],
        ["indent_feedback",     "Stores user / manager feedback on findings, used offline to improve the service"],
        ["validation_standard", "Stores the versioned best-practice standard per domain, read by the service at startup"],
    ],
)

ANALYSIS_TBL = (
    ["Column", "Type", "Notes"],
    [
        ["analysis_id",              "string",        "Primary key"],
        ["indent_id",                "string",        "Indexed; links to the indent"],
        ["domain",                   "string",        "Indent domain"],
        ["analysis_timestamp",       "datetime",      "When analysis ran"],
        ["score",                    "number",        "0-100"],
        ["grade",                    "string",        "Strong / Adequate / Needs Improvement / Weak"],
        ["score_breakdown",          "object (JSON)", "Per-category points"],
        ["findings",                 "list (JSON)",   "Detailed findings"],
        ["recommendations",          "list (JSON)",   "Recommended actions"],
        ["missing_documents",        "list (JSON)",   "Required-but-absent document types"],
        ["interrelationship_issues", "list (JSON)",   "Cross-document inconsistencies"],
        ["document_inventory",       "list (JSON)",   "Classified documents + genuine flags"],
        ["standard_version",         "string",        "Which standard version was used"],
        ["is_latest",                "boolean",       "True for the most recent per indent"],
        ["full_response",            "object (JSON)", "Complete response, for audit"],
    ],
)

FEEDBACK_TBL = (
    ["Column", "Type", "Notes"],
    [
        ["feedback_id",        "string",   "Primary key"],
        ["analysis_id",        "string",   "Links to indent_analysis"],
        ["indent_id",          "string",   "Indexed"],
        ["domain",             "string",   "Indent domain"],
        ["feedback_timestamp", "datetime", "When feedback was given"],
        ["feedback_type",      "string",   "finding_wrong / finding_correct / severity_wrong / missing_finding"],
        ["finding_category",   "string",   "mandatory / documentation / risk / vendor / approval"],
        ["finding_title",      "string",   "The finding being commented on"],
        ["user_comment",       "text",     "Free-text explanation"],
        ["correct_value",      "text",     "What the user says is correct (optional)"],
        ["submitted_by",       "string",   "User or manager identifier"],
    ],
)

STANDARD_TBL = (
    ["Column", "Type", "Notes"],
    [
        ["standard_id",         "string",        "Primary key"],
        ["domain",              "string",        "Indent domain this standard applies to"],
        ["version",             "string",        "Version label (e.g. v4-2026-06)"],
        ["published_timestamp", "datetime",      "When this version was published"],
        ["standard_body",       "object (JSON)", "The full standard content"],
        ["is_current",          "boolean",       "True for the active version per domain"],
    ],
)

SYSTEMS_TBL = (
    ["System", "Role", "Integration"],
    [
        ["IPMS Portal",    "Triggers analysis, displays results, owns the database", "Calls the Check-Adequacy API; persists the response"],
        ["IPMS Field API", "Source of indent field data",                            "Service fetches question / answer / required entries"],
        ["Document Store", "Source of attachments",                                  "Service fetches files via authenticated OAuth"],
        ["GENAI API",      "Extraction and analysis",                                "Service sends prepared content over an authenticated session"],
    ],
)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    def heading(text, size, bold=True, space_before=12, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Arial"
        return p

    def para(text, bold_lead=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if bold_lead:
            r = p.add_run(bold_lead)
            r.bold = True
        p.add_run(text)
        return p

    def bullet(text, bold_lead=None):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if bold_lead:
            r = p.add_run(bold_lead)
            r.bold = True
        p.add_run(text)
        return p

    def mono_block(text):
        # render an ASCII block in Courier New so it lines up
        for line in text.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(line if line else " ")
            r.font.name = "Courier New"
            r.font.size = Pt(8)

    def add_table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = "Arial"
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(val)
                run.font.size = Pt(9.5)
                run.font.name = "Arial"
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return t

    def flow_box(title, lines):
        # one-cell bordered table per step
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        cell = t.rows[0].cells[0]
        cell.text = ""
        rt = cell.paragraphs[0].add_run(title)
        rt.bold = True
        rt.font.size = Pt(10)
        rt.font.name = "Arial"
        for ln in lines:
            pp = cell.add_paragraph()
            rr = pp.add_run(ln)
            rr.font.size = Pt(9)
            rr.font.name = "Arial"
        # arrow paragraph
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = ap.add_run("\u2193")
        ar.font.size = Pt(12)
        ar.bold = True

    # ---- TITLE ----
    tp = doc.add_paragraph()
    tr = tp.add_run(TITLE)
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.name = "Arial"
    sp = doc.add_paragraph()
    sr = sp.add_run(SUBTITLE)
    sr.font.size = Pt(12)
    vp = doc.add_paragraph()
    vr = vp.add_run(VERSION)
    vr.font.name = "Courier New"
    vr.font.size = Pt(9)

    # ---- 1. EXEC SUMMARY ----
    heading("1.  Executive Summary", 15)
    para("The Indent Validation Service is a stateless service that integrates with IPMS to evaluate procurement indents and return an adequacy score together with a detailed analysis. It is invoked when a user selects the \u201CCheck Adequacy\u201D action on an indent in IPMS.")
    para("This document focuses on production integration: the APIs the service exposes and consumes, the database tables IPMS must create, the system architecture, and the runtime data flow. The internal analysis logic is summarised only briefly, as it is not the integration concern.")
    para("The service does not store data. It receives an indent identifier, gathers the indent's field data and document attachments, analyses them, and returns a structured response. IPMS persists the response - the score and the full detailed analysis - and manages all history and display.")
    para("The adequacy score is guidance, not a gate. Users may submit at any score. Both the user and the manager see the same detailed analysis.", bold_lead="Advisory, not blocking.  ")
    para("The system has two subsystems:")
    bullet(" - handles each request, returns results, reads the standard from memory, never writes to the database.", bold_lead="Live Validation API")
    bullet(" - scheduled; rebuilds the standard from accumulated indents and publishes a new version to the database.", bold_lead="Offline Standard Pipeline")
    para("All language-model calls are made through an authenticated GENAI API over an authenticated session.", bold_lead="LLM access.  ")

    # ---- 2. API CONTRACTS ----
    heading("2.  API Contracts", 15)
    para("This section defines every integration point. Each contract is described as a field table: the key name, its data type (string, number, boolean, list, or object), and its meaning.")
    heading("APIs required", 11)
    para("The integration involves one API the service exposes and three it consumes:")
    add_table(*APIS_REQUIRED)

    heading("2.1  Check-Adequacy API - request", 12)
    para("POST /check-adequacy  - body is an object:")
    add_table(*REQ_REQUEST)
    para("Immediate response (before analysis completes) - an object:")
    add_table(*REQ_ACK)

    heading("2.2  Check-Adequacy API - result", 12)
    para("Delivered when analysis completes. This is the full payload IPMS persists - the score and the complete detailed analysis. An object with these keys:")
    add_table(*RESULT_TBL)
    para("IPMS persists the entire detailed analysis alongside the score, so the user and manager always see the full reasoning, not just the number.", bold_lead="Detailed analysis stored with the score.  ")

    heading("2.3  Indent Field API - consumed", 12)
    para("Returns an object with indent_id and fields, where fields is a list of objects - one per filled field:")
    add_table(*FIELD_API_TBL)
    para("Fields are conditional - selecting one option unlocks others - so the key set varies between indents. The service consumes whatever question / answer / required entries are returned; new or renamed fields do not break it.", bold_lead="Dynamic fields.  ")

    heading("2.4  Document Store API - consumed", 12)
    para("Returns the indent's attachments over an authenticated (OAuth) connection: a list of files (binary content plus original filename). Filenames are arbitrary and are not relied upon - each file's true type is determined internally.")

    # ---- 3. DATABASE ----
    heading("3.  Database Design", 15)
    para("The Validation Service does not write to the database. IPMS persists the service's response and manages history. This section defines the tables IPMS must create.")
    heading("Tables required", 11)
    add_table(*TABLES_REQUIRED)
    heading("3.1  indent_analysis", 12)
    para("Stores the score and the full detailed analysis. History is preserved by keeping all rows; the latest per indent is flagged for display.")
    add_table(*ANALYSIS_TBL)
    para("On a new analysis for an indent, set is_latest = false on all prior rows for that indent_id, then insert the new row with is_latest = true. Display filters on is_latest = true; history returns all rows by timestamp.", bold_lead="How \u201Clatest only\u201D works.  ")
    heading("3.2  indent_feedback", 12)
    para("Captures user / manager feedback on findings - read offline to improve prompts and correct the standard. Each row links to an analysis.")
    add_table(*FEEDBACK_TBL)
    heading("3.3  validation_standard", 12)
    para("Stores the versioned best-practice standard per domain. The service loads the current version into memory at startup and reloads when a new version is published.")
    add_table(*STANDARD_TBL)

    # ---- 4. ARCHITECTURE ----
    heading("4.  High-Level Architecture", 15)
    para("The diagram below shows the systems involved and how they connect. The Validation Service is stateless except for the in-memory standard. Solid flow is runtime; the offline pipeline publishes the standard separately.")
    mono_block(ARCH_DIAGRAM)
    doc.add_paragraph()
    para("Four systems participate:")
    add_table(*SYSTEMS_TBL)
    para("The service holds no persistent indent state. Each request is self-contained: fetch, analyse, return, forget. The only in-memory state is the current standard, loaded at startup and refreshed on a new version.", bold_lead="Statelessness.  ")

    # ---- 5. DATA FLOW ----
    heading("5.  Data Flow - \u201CCheck Adequacy\u201D", 15)
    para("This is the runtime path, triggered each time a user selects \u201CCheck Adequacy\u201D. It is asynchronous: the service acknowledges the request immediately and delivers the result when analysis completes.")
    for title, lines in FLOW_STEPS:
        flow_box(title, lines)
    heading("5.1  Internal Analysis Steps (Summary)", 12)
    para("Steps 4 and 5 are performed inside the service and are not integration concerns. In brief:")
    bullet(" - attachments are identified by what they contain, not by filename or upload slot, and verified as genuine.", bold_lead="Classify documents by content")
    bullet(" - the service derives what the indent should contain from the field selections, then checks the actual content against that.", bold_lead="Infer requirements")

    # ---- 6. STANDARD ----
    heading("6.  Standard Storage & Refresh", 15)
    para("The best-practice standard is the benchmark every indent is measured against. It is stored in the validation_standard table and its lifecycle is separated from live traffic.")
    bullet("Stored as a versioned record per domain (validation_standard table).")
    bullet("Small and changes rarely.")
    bullet("The live service loads the current version into memory at startup.")
    bullet("The service periodically checks for a newer version and reloads if found - a rebuild takes effect without redeployment.")
    para("Because the standard lives in memory, comparison is instant. The database is touched only at startup and on version change - never on the critical path of a request.", bold_lead="No per-request database access for the standard.  ")
    para("The offline pipeline rebuilds the standard on a schedule (a hybrid trigger: rebuild when 15 new quality-passed indents accumulate, or 30 days pass, whichever is first), then publishes a new version to the validation_standard table. Each analysis records the standard_version it used, so past scores remain reproducible.")

    doc.save(path)
    print("DOCX written:", path)


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def build_pdf(path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, Preformatted)
    from reportlab.lib.enums import TA_CENTER

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                          fontSize=9.5, leading=13, spaceAfter=5)
    h1 = ParagraphStyle("h1", parent=styles["Normal"], fontName="Helvetica-Bold",
                        fontSize=15, leading=18, spaceBefore=14, spaceAfter=7)
    h2 = ParagraphStyle("h2", parent=styles["Normal"], fontName="Helvetica-Bold",
                        fontSize=12, leading=15, spaceBefore=10, spaceAfter=5)
    h3 = ParagraphStyle("h3", parent=styles["Normal"], fontName="Helvetica-Bold",
                        fontSize=11, leading=14, spaceBefore=8, spaceAfter=4)
    titlest = ParagraphStyle("title", parent=styles["Normal"], fontName="Helvetica-Bold",
                             fontSize=20, leading=24, spaceAfter=2)
    subst = ParagraphStyle("sub", parent=styles["Normal"], fontName="Helvetica",
                           fontSize=12, leading=15, spaceAfter=2)
    verst = ParagraphStyle("ver", parent=styles["Normal"], fontName="Courier",
                           fontSize=9, leading=12, spaceAfter=10)
    monost = ParagraphStyle("mono", parent=styles["Normal"], fontName="Courier",
                            fontSize=7, leading=8.4)
    boxtitle = ParagraphStyle("bt", parent=styles["Normal"], fontName="Helvetica-Bold",
                              fontSize=9.5, leading=12)
    boxline = ParagraphStyle("bl", parent=styles["Normal"], fontName="Helvetica",
                             fontSize=8.5, leading=11)
    arrowst = ParagraphStyle("ar", parent=styles["Normal"], fontName="Helvetica-Bold",
                             fontSize=12, leading=14, alignment=TA_CENTER, spaceBefore=2, spaceAfter=2)

    elements = []
    PAGE_W = letter[0] - 1.5 * inch   # usable width with 0.75" margins

    def esc(t):
        return (t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def P(text, st=body, lead=None):
        if lead:
            text = "<b>" + esc(lead) + "</b>" + esc(text)
        else:
            text = esc(text)
        elements.append(Paragraph(text, st))

    def bullet_pdf(text, lead=None):
        inner = ("<b>" + esc(lead) + "</b>" + esc(text)) if lead else esc(text)
        elements.append(Paragraph("\u2022&nbsp;&nbsp;" + inner,
                        ParagraphStyle("blt", parent=body, leftIndent=14, spaceAfter=3)))

    def table_pdf(headers, rows, rel_widths):
        total = sum(rel_widths)
        col_w = [PAGE_W * w / total for w in rel_widths]
        data = [[Paragraph("<b>" + esc(h) + "</b>",
                 ParagraphStyle("th", parent=body, textColor=colors.white, fontSize=9))
                 for h in headers]]
        for row in rows:
            data.append([Paragraph(esc(c), ParagraphStyle("td", parent=body, fontSize=8.5, leading=11))
                         for c in row])
        t = Table(data, colWidths=col_w, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F1F3D")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#888888")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F5F7")]),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 6))

    def flow_box_pdf(title, lines):
        cell = [Paragraph(esc(title), boxtitle)]
        for ln in lines:
            cell.append(Paragraph(esc(ln), boxline))
        t = Table([[cell]], colWidths=[PAGE_W * 0.82])
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 1, colors.black),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        elements.append(t)
        elements.append(Paragraph("\u2193", arrowst))

    # TITLE
    P(TITLE, titlest)
    P(SUBTITLE, subst)
    P(VERSION, verst)

    # 1
    P("1.  Executive Summary", h1)
    P("The Indent Validation Service is a stateless service that integrates with IPMS to evaluate procurement indents and return an adequacy score together with a detailed analysis. It is invoked when a user selects the \u201CCheck Adequacy\u201D action on an indent in IPMS.")
    P("This document focuses on production integration: the APIs the service exposes and consumes, the database tables IPMS must create, the system architecture, and the runtime data flow. The internal analysis logic is summarised only briefly, as it is not the integration concern.")
    P("The service does not store data. It receives an indent identifier, gathers the indent's field data and document attachments, analyses them, and returns a structured response. IPMS persists the response - the score and the full detailed analysis - and manages all history and display.")
    P("The adequacy score is guidance, not a gate. Users may submit at any score. Both the user and the manager see the same detailed analysis.", lead="Advisory, not blocking.  ")
    P("The system has two subsystems:")
    bullet_pdf(" - handles each request, returns results, reads the standard from memory, never writes to the database.", lead="Live Validation API")
    bullet_pdf(" - scheduled; rebuilds the standard from accumulated indents and publishes a new version to the database.", lead="Offline Standard Pipeline")
    P("All language-model calls are made through an authenticated GENAI API over an authenticated session.", lead="LLM access.  ")

    # 2
    P("2.  API Contracts", h1)
    P("This section defines every integration point. Each contract is described as a field table: the key name, its data type (string, number, boolean, list, or object), and its meaning.")
    P("APIs required", h3)
    P("The integration involves one API the service exposes and three it consumes:")
    table_pdf(*APIS_REQUIRED, [2.2, 3.0, 4.0])
    P("2.1  Check-Adequacy API - request", h2)
    P("POST /check-adequacy  - body is an object:")
    table_pdf(*REQ_REQUEST, [2.2, 1.4, 5.4])
    P("Immediate response (before analysis completes) - an object:")
    table_pdf(*REQ_ACK, [2.2, 1.4, 5.4])
    P("2.2  Check-Adequacy API - result", h2)
    P("Delivered when analysis completes. This is the full payload IPMS persists - the score and the complete detailed analysis. An object with these keys:")
    table_pdf(*RESULT_TBL, [2.6, 1.8, 4.6])
    P("IPMS persists the entire detailed analysis alongside the score, so the user and manager always see the full reasoning, not just the number.", lead="Detailed analysis stored with the score.  ")
    P("2.3  Indent Field API - consumed", h2)
    P("Returns an object with indent_id and fields, where fields is a list of objects - one per filled field:")
    table_pdf(*FIELD_API_TBL, [3.0, 1.4, 4.6])
    P("Fields are conditional - selecting one option unlocks others - so the key set varies between indents. The service consumes whatever question / answer / required entries are returned; new or renamed fields do not break it.", lead="Dynamic fields.  ")
    P("2.4  Document Store API - consumed", h2)
    P("Returns the indent's attachments over an authenticated (OAuth) connection: a list of files (binary content plus original filename). Filenames are arbitrary and are not relied upon - each file's true type is determined internally.")

    # 3
    P("3.  Database Design", h1)
    P("The Validation Service does not write to the database. IPMS persists the service's response and manages history. This section defines the tables IPMS must create.")
    P("Tables required", h3)
    table_pdf(*TABLES_REQUIRED, [2.6, 6.4])
    P("3.1  indent_analysis", h2)
    P("Stores the score and the full detailed analysis. History is preserved by keeping all rows; the latest per indent is flagged for display.")
    table_pdf(*ANALYSIS_TBL, [2.8, 1.8, 4.4])
    P("On a new analysis for an indent, set is_latest = false on all prior rows for that indent_id, then insert the new row with is_latest = true. Display filters on is_latest = true; history returns all rows by timestamp.", lead="How \u201Clatest only\u201D works.  ")
    P("3.2  indent_feedback", h2)
    P("Captures user / manager feedback on findings - read offline to improve prompts and correct the standard. Each row links to an analysis.")
    table_pdf(*FEEDBACK_TBL, [2.6, 1.5, 4.9])
    P("3.3  validation_standard", h2)
    P("Stores the versioned best-practice standard per domain. The service loads the current version into memory at startup and reloads when a new version is published.")
    table_pdf(*STANDARD_TBL, [2.8, 1.8, 4.4])

    # 4
    P("4.  High-Level Architecture", h1)
    P("The diagram below shows the systems involved and how they connect. The Validation Service is stateless except for the in-memory standard. Solid flow is runtime; the offline pipeline publishes the standard separately.")
    elements.append(Preformatted(ARCH_DIAGRAM.strip("\n"), monost))
    elements.append(Spacer(1, 8))
    P("Four systems participate:")
    table_pdf(*SYSTEMS_TBL, [1.9, 3.6, 3.5])
    P("The service holds no persistent indent state. Each request is self-contained: fetch, analyse, return, forget. The only in-memory state is the current standard, loaded at startup and refreshed on a new version.", lead="Statelessness.  ")

    # 5
    P("5.  Data Flow - \u201CCheck Adequacy\u201D", h1)
    P("This is the runtime path, triggered each time a user selects \u201CCheck Adequacy\u201D. It is asynchronous: the service acknowledges the request immediately and delivers the result when analysis completes.")
    for title, lines in FLOW_STEPS:
        flow_box_pdf(title, lines)
    P("5.1  Internal Analysis Steps (Summary)", h2)
    P("Steps 4 and 5 are performed inside the service and are not integration concerns. In brief:")
    bullet_pdf(" - attachments are identified by what they contain, not by filename or upload slot, and verified as genuine.", lead="Classify documents by content")
    bullet_pdf(" - the service derives what the indent should contain from the field selections, then checks the actual content against that.", lead="Infer requirements")

    # 6
    P("6.  Standard Storage & Refresh", h1)
    P("The best-practice standard is the benchmark every indent is measured against. It is stored in the validation_standard table and its lifecycle is separated from live traffic.")
    bullet_pdf("Stored as a versioned record per domain (validation_standard table).")
    bullet_pdf("Small and changes rarely.")
    bullet_pdf("The live service loads the current version into memory at startup.")
    bullet_pdf("The service periodically checks for a newer version and reloads if found - a rebuild takes effect without redeployment.")
    P("Because the standard lives in memory, comparison is instant. The database is touched only at startup and on version change - never on the critical path of a request.", lead="No per-request database access for the standard.  ")
    P("The offline pipeline rebuilds the standard on a schedule (a hybrid trigger: rebuild when 15 new quality-passed indents accumulate, or 30 days pass, whichever is first), then publishes a new version to the validation_standard table. Each analysis records the standard_version it used, so past scores remain reproducible.")

    doc = SimpleDocTemplate(path, pagesize=letter,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                            topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    doc.build(elements)
    print("PDF written:", path)


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(here, "Indent_Validation_Service_Architecture.docx")
    pdf_path  = os.path.join(here, "Indent_Validation_Service_Architecture.pdf")

    try:
        build_docx(docx_path)
    except Exception as e:
        print("DOCX generation failed:", e)

    try:
        build_pdf(pdf_path)
    except Exception as e:
        print("PDF generation failed:", e)

    print("\nDone. Files are in:", here)
