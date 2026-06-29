import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

print("Tesseract path:", pytesseract.pytesseract.tesseract_cmd)

try:
    print("Tesseract version:", pytesseract.get_tesseract_version())
except Exception as e:
    print("Tesseract ERROR:", e)

import os
import io
from pathlib import Path
from typing import Dict, Any

import pandas as pd
import fitz
from PIL import Image
from docx import Document

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".xlsx",
    ".xlsm",
    ".xls",
    ".docx",
    ".txt",
}


def should_use_ocr(extracted_text: str, min_chars: int = 500) -> bool:
    if not extracted_text or not extracted_text.strip():
        return True

    cleaned = extracted_text.strip()

    if len(cleaned) < min_chars:
        return True

    alpha_count = sum(ch.isalpha() for ch in cleaned)
    alpha_ratio = alpha_count / max(len(cleaned), 1)

    return alpha_ratio < 0.25


def parse_pdf_text(pdf_path: Path) -> tuple[str, Dict[str, Any]]:
    text_parts = []
    doc = fitz.open(pdf_path)

    for page_index, page in enumerate(doc):
        page_text = page.get_text("text") or ""
        text_parts.append(f"\n--- Page {page_index + 1} ---\n{page_text}")

    text = "\n".join(text_parts)

    metadata = {
        "parser":   "pymupdf",
        "pages":    len(doc),
        "ocr_used": False,
    }

    use_ocr_env = os.getenv("USE_OCR", "true").lower() == "true"

    if use_ocr_env and should_use_ocr(text):
        ocr_text_parts = []

        for page_index, page in enumerate(doc):
            pix          = page.get_pixmap(dpi=200)
            image_bytes  = pix.tobytes("png")
            image        = Image.open(io.BytesIO(image_bytes))

            page_ocr_text = pytesseract.image_to_string(image)
            ocr_text_parts.append(f"\n--- Page {page_index + 1} OCR ---\n{page_ocr_text}")

        text                   = "\n".join(ocr_text_parts)
        metadata["parser"]     = "pymupdf+tesseract"
        metadata["ocr_used"]   = True

    return text, metadata


def parse_excel_text(excel_path: Path) -> tuple[str, Dict[str, Any]]:
    suffix = excel_path.suffix.lower()

    engine = "openpyxl"
    if suffix == ".xls":
        engine = "xlrd"

    sheets = pd.read_excel(
        excel_path,
        sheet_name=None,
        dtype=str,
        engine=engine,
    )

    sheet_metadata = {}

    def score_sheet(df: "pd.DataFrame") -> float:
        df = df.fillna("")
        total_cells = max(df.size, 1)

        # Non-empty cells
        non_empty = sum(1 for v in df.values.flatten() if str(v).strip())
        fill_ratio = non_empty / total_cells

        # Numeric cell density — diagram sheets have almost none
        numeric_cells = sum(
            1 for v in df.values.flatten()
            if str(v).strip().replace(".", "").replace(",", "").replace("-", "").isdigit()
            and str(v).strip() not in ("", "-")
        )
        number_density = numeric_cells / total_cells

        # Row count — data sheets have more rows
        row_score = min(df.shape[0] / 20, 5)

        # Generic procurement keyword bonus
        all_text = " ".join(
            str(v).lower() for v in df.values.flatten() if str(v).strip()
        )
        keywords = [
            "quantity", "qty", "rate", "amount", "description",
            "total", "unit", "cost", "price", "item", "sl",
            "scope", "material", "labour", "service", "vendor",
        ]
        keyword_score = sum(1 for k in keywords if k in all_text)

        return (fill_ratio * 5) + (number_density * 15) + row_score + keyword_score

    # Score all sheets, pick the best
    best_sheet_name = None
    best_score      = -1

    for sheet_name, df in sheets.items():
        df = df.fillna("")
        sheet_metadata[sheet_name] = {
            "rows":    int(df.shape[0]),
            "columns": int(df.shape[1]),
            "headers": [str(c) for c in df.columns],
        }
        s = score_sheet(df)
        sheet_metadata[sheet_name]["data_score"] = round(s, 2)
        if s > best_score:
            best_score      = s
            best_sheet_name = sheet_name

    # Extract text from best sheet only
    text_parts = []
    if best_sheet_name and best_sheet_name in sheets:
        df = sheets[best_sheet_name].fillna("")
        text_parts.append(
            f"\n=== Sheet: {best_sheet_name} "
            f"(selected from {len(sheets)} sheets, score={best_score:.1f}) ==="
        )
        text_parts.append("Headers: " + " | ".join([str(c) for c in df.columns]))

        for row_index, row in df.iterrows():
            row_values = []
            for col in df.columns:
                value = str(row[col]).strip()
                if value:
                    row_values.append(f"{col}: {value}")
            if row_values:
                text_parts.append(
                    f"Row {row_index + 1}: " + " | ".join(row_values)
                )

    other_sheets = [n for n in sheets if n != best_sheet_name]
    if other_sheets:
        text_parts.append(
            f"\n[Skipped sheets (low data score): {', '.join(other_sheets)}]"
        )

    metadata = {
        "parser":           "pandas.read_excel",
        "sheets":           sheet_metadata,
        "best_sheet":       best_sheet_name,
        "best_sheet_score": round(best_score, 2),
        "ocr_used":         False,
    }

    return "\n".join(text_parts), metadata


def parse_docx_text(docx_path: Path) -> tuple[str, Dict[str, Any]]:
    document  = Document(docx_path)
    text_parts = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    for table_index, table in enumerate(document.tables):
        text_parts.append(f"\n=== Table {table_index + 1} ===")
        for row_index, row in enumerate(table.rows):
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            text_parts.append(f"Row {row_index + 1}: " + " | ".join(values))

    metadata = {
        "parser":     "python-docx",
        "paragraphs": len(document.paragraphs),
        "tables":     len(document.tables),
        "ocr_used":   False,
    }

    return "\n".join(text_parts), metadata


def parse_txt_text(txt_path: Path) -> tuple[str, Dict[str, Any]]:
    text = txt_path.read_text(encoding="utf-8", errors="ignore")

    metadata = {
        "parser":   "plain_text",
        "ocr_used": False,
    }

    return text, metadata


def parse_file(file_path: str | Path) -> tuple[str, Dict[str, Any]]:
    file_path = Path(file_path)
    suffix    = file_path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf_text(file_path)

    if suffix in {".xlsx", ".xlsm", ".xls"}:
        return parse_excel_text(file_path)

    if suffix == ".docx":
        return parse_docx_text(file_path)

    if suffix == ".txt":
        return parse_txt_text(file_path)

    raise ValueError(f"Unsupported file type: {file_path}")

